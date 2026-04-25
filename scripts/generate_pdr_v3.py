from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = "canteen_pdr_v3.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12, "1F4E79"),
        ("Title", 24, "17365D"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(10.5)
        styles[name].paragraph_format.space_after = Pt(3)


def para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def h1(doc, n, title):
    doc.add_heading(f"{n}. {title}", level=1)


def h2(doc, n, title):
    doc.add_heading(f"{n} {title}", level=2)


def kv(doc, pairs):
    for key, value in pairs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{key}: ")
        r.bold = True
        p.add_run(value)


def build_doc():
    doc = Document()
    setup(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PRODUCT DESIGN REPORT")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(23, 54, 93)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("College Canteen Management System")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)

    kv(
        doc,
        [
            ("Document version", "v3.0 - Production Design Baseline"),
            ("Status", "Ready for engineering planning"),
            ("Stack", "Next.js 14 App Router, Firebase, Razorpay, Vercel"),
            ("Audience", "Engineering, QA, DevOps, canteen operations"),
            ("Business timezone", "Asia/Kolkata (IST)"),
            ("Date", "April 2026"),
        ],
    )

    para(
        doc,
        "Design principle: All money, identity, role, order-state, availability, and refund decisions are enforced on the server. The browser is a user interface, not a source of authority.",
        "Design principle:",
    )

    h1(doc, 1, "Executive Summary")
    para(
        doc,
        "The College Canteen Management System is a mobile-first ordering and kitchen operations platform for a single college canteen. It lets students browse the menu, pay online, receive a daily pickup token, track preparation status, and collect food when ready. Kitchen staff get a live queue. Administrators manage menu, staff, availability, reporting, and operational controls.",
    )
    para(
        doc,
        "This v3 design corrects the production gaps in the earlier draft: first-admin bootstrap, payment failure recovery, webhook idempotency, explicit IST handling, authenticated ISR revalidation, historical stats retention, rate limits, cart constraints, storage rules, testing, and operations.",
    )

    h1(doc, 2, "Problem Statement")
    bullets(
        doc,
        [
            "Students lose time in physical queues because order placement, payment, and preparation are coupled at the counter.",
            "Kitchen staff lack a structured queue, causing missed, duplicated, or verbally miscommunicated orders.",
            "Unavailable items are discovered too late, creating wasted trips and substitutions.",
            "Manual payments reduce throughput and create reconciliation risk.",
            "The canteen has no reliable order history, payment trail, demand history, or daily revenue audit.",
        ],
    )

    h1(doc, 3, "Goals and Non-Goals")
    h2(doc, "3.1", "Launch Goals")
    bullets(
        doc,
        [
            "Student ordering: browse menu, pay, receive token, track order, and cancel within the defined window.",
            "Kitchen operations: live queue with status transitions, elapsed time, item details, and prominent offline warnings.",
            "Administration: menu CRUD, staff account management, availability controls, daily and historical reporting.",
            "Payment integrity: no confirmed Firestore order exists until payment is verified by signature or trusted webhook.",
            "Scale target: 200 concurrent users and 50 active kitchen orders without visible degradation.",
            "Auditability: immutable payment, refund, status, and daily stats records.",
        ],
    )
    h2(doc, "3.2", "MVP Non-Goals")
    bullets(
        doc,
        [
            "Multi-canteen support.",
            "Per-item inventory stock counting.",
            "Delivery to tables or hostels.",
            "Loyalty, coupons, recommendations, or AI forecasting.",
        ],
    )

    h1(doc, 4, "Users and Access Model")
    bullets(
        doc,
        [
            "Student: Firebase Phone Auth OTP. Capabilities: menu, cart, payment, tracking, eligible cancellation, order history.",
            "Staff: Firebase email/password plus staff custom claim. Capabilities: live queue, status transitions, item availability toggles.",
            "Admin: Firebase email/password plus admin custom claim and MFA. Capabilities: menu CRUD, staff provisioning, reports, configuration, audit access.",
        ],
    )
    h2(doc, "4.1", "First-Admin Bootstrap")
    para(
        doc,
        "The first admin cannot be created by assignRole because assignRole itself requires an existing admin. The first admin is created once using a local-only Node.js script at scripts/bootstrap-admin.js.",
    )
    numbers(
        doc,
        [
            "Create the first admin user in Firebase Auth manually or through a local script.",
            "Run scripts/bootstrap-admin.js from a trusted developer machine using Firebase Admin SDK credentials.",
            "Set the custom claim { role: \"admin\" } on the designated UID and write users/{uid}.roleDisplay = \"admin\".",
            "Never deploy this script to Cloud Functions. Document the procedure in README deployment steps.",
            "After bootstrap, all role changes happen only through assignRole and are audit logged.",
        ],
    )

    h1(doc, 5, "System Architecture")
    bullets(
        doc,
        [
            "Frontend: Next.js 14 App Router for student, staff, and admin interfaces.",
            "Hosting/CDN: Vercel for static/ISR delivery, preview deployments, and production domain.",
            "Authentication: Firebase Authentication for OTP, email/password, and custom claims.",
            "Database: Cloud Firestore for users, menu items, orders, payment attempts, counters, stats, and rate limits.",
            "Server logic: Firebase Cloud Functions v2 on Node 20 for all privileged mutations.",
            "Payment: Razorpay for payment orders, payment sheet, signatures, webhooks, and refunds.",
            "Storage: Firebase Storage for menu images and generated WebP assets.",
            "Integrity: Firebase App Check enforced in staging and production.",
            "Monitoring: Google Cloud/Firebase logs and alerts for payments, webhooks, refunds, latency, and queue health.",
        ],
    )

    h1(doc, 6, "Environment and Secrets")
    bullets(
        doc,
        [
            "Firebase projects: canteen-dev, canteen-staging, canteen-prod.",
            "Razorpay keys: test keys in development/staging; live keys only in production Secret Manager.",
            "Firestore rules: emulator-only relaxation in development; production rules in staging and production.",
            "App Check: disabled/debug in development, debug token in staging, reCAPTCHA v3 enforced in production.",
            "ISR revalidation: REVALIDATE_SECRET is stored in Vercel environment variables and Cloud Function Secret Manager.",
            "Allowed origins: localhost for development, Vercel preview domains for staging, production domain for production.",
            "Secrets must never be committed, logged, returned to clients, or placed in browser-visible environment variables.",
            "Next.js route handlers validate REVALIDATE_SECRET before calling revalidatePath or revalidateTag.",
        ],
    )

    h1(doc, 7, "Timezone and Business Calendar")
    bullets(
        doc,
        [
            "All business-day concepts use Asia/Kolkata. Firestore timestamps remain UTC instants.",
            "Scheduled functions use timeZone: \"Asia/Kolkata\" and a schedule such as \"every day 00:00\".",
            "Daily stats document IDs use the IST local date: stats/YYYY-MM-DD.",
            "Hourly stats keys use IST hours from 00 through 23.",
            "Order display times are converted to IST for staff/admin UI.",
            "Tests must include orders around 23:59 IST and 00:00 IST.",
        ],
    )

    h1(doc, 8, "Data Model")
    h2(doc, "8.1", "Core Collections")
    bullets(
        doc,
        [
            "users/{uid}: profile and display metadata. Role display is not authoritative; custom claims are authoritative.",
            "menuItems/{itemId}: menu catalog. Price is stored as integer paisa; availability is checked server-side.",
            "orders/{razorpayOrderId}: confirmed paid orders. The deterministic document ID supports idempotency.",
            "paymentAttempts/{razorpayOrderId}: tracks initiated, failed, expired, pending, and captured payment attempts.",
            "stats/today: live dashboard alias updated in real time.",
            "stats/YYYY-MM-DD: historical daily stats. Historical documents are never overwritten after finalisation except audited correction.",
            "counters/daily: current token counter, updated only inside payment confirmation transaction.",
            "rateLimits/{uid}: per-user function counters for throttling.",
        ],
    )
    h2(doc, "8.2", "Order Fields")
    bullets(
        doc,
        [
            "orderId: same as razorpayOrderId.",
            "token: daily pickup token generated transactionally in the IST business day.",
            "userId and customerName: UID plus display snapshot at order creation.",
            "items: snapshot of itemId, name, pricePaisa, and quantity.",
            "totalPaisa: computed server-side and never accepted from client.",
            "status: placed, preparing, ready, completed, cancelled.",
            "paymentStatus: captured, refunded, refund_pending, refund_failed.",
            "refund: refundId, status, speed, createdAt, processedAt, and failureReason when present.",
            "statusHistory: append-only list of state changes with actor and timestamp.",
        ],
    )

    h1(doc, 9, "Payment Lifecycle")
    para(
        doc,
        "Invariant: A confirmed Firestore order is created only after Razorpay payment success is verified by HMAC signature or by a trusted signed webhook. initiatePayment creates a Razorpay order and a paymentAttempts document, not a confirmed order.",
        "Invariant:",
    )
    numbers(
        doc,
        [
            "Student taps Pay. Client calls initiatePayment with cart payload.",
            "Cloud Function validates auth, App Check, rate limit, active order limit, item availability, quantities, and total.",
            "Function creates a Razorpay order object and writes paymentAttempts/{razorpayOrderId}.",
            "Razorpay payment sheet handles UPI, card, wallet, or other payment method.",
            "On success, client calls verifyPayment with payment_id, order_id, and signature.",
            "Razorpay may also send payment.captured webhook. The handler verifies the webhook signature.",
            "Both paths call the same createConfirmedOrder routine using Firestore create() on orders/{razorpayOrderId}.",
            "If the document already exists, the second caller reads and returns the existing order.",
            "Client shows token and starts a single-order listener.",
        ],
    )
    h2(doc, "9.1", "Payment Failures and Recovery")
    bullets(
        doc,
        [
            "Razorpay API unavailable during initiatePayment: return a retry message; no order is created.",
            "Student payment fails in Razorpay sheet: mark paymentAttempts as failed if callback is received; show retry option.",
            "Network drops after payment capture: show confirming-payment screen; webhook can create the order asynchronously.",
            "verifyPayment signature mismatch: log identifiers, return a generic error, and never create order.",
            "Webhook retry: handler remains idempotent and safe to run multiple times.",
            "Refund API failure: mark refund_pending or refund_failed and alert admin.",
        ],
    )

    h1(doc, 10, "Business Rules")
    bullets(
        doc,
        [
            "Max active orders per student: one placed/preparing/ready order.",
            "Max distinct items: 10.",
            "Max quantity per item: 5.",
            "Max order value: configured by canteen; default INR 2,000.",
            "Cancellation: allowed only while status is placed and within 90 seconds of createdAt.",
            "Unavailable item: rejected for new orders; existing confirmed orders are unaffected.",
            "Token reset: daily at 00:00 Asia/Kolkata.",
        ],
    )

    h1(doc, 11, "Order State Machine")
    bullets(
        doc,
        [
            "payment verified -> placed: verifyPayment/webhook only.",
            "placed -> preparing: staff only.",
            "placed -> cancelled: student/admin only, within cancellation window, with refund initiation.",
            "preparing -> ready: staff only.",
            "ready -> completed: staff only.",
            "preparing -> cancelled is not permitted for student cancellation.",
            "Status transitions should go through Cloud Functions to centralise validation, history, and audit logging.",
            "Firestore rules still reject invalid direct client writes as a backstop.",
        ],
    )

    h1(doc, 12, "Security Design")
    bullets(
        doc,
        [
            "Custom claims are the only source of role authority.",
            "App Check is enforced in staging and production, but it is not the only security control.",
            "Every callable validates request.auth, App Check, role claim, payload shape, and business rules.",
            "Callable and HTTP functions use explicit CORS allowlists.",
            "Razorpay signatures are verified using HMAC-SHA256. Secrets are never browser-visible.",
            "Security rules deny direct order creation and counter writes from clients.",
            "All role changes, refunds, and status changes are audit logged.",
        ],
    )

    h1(doc, 13, "Firestore Rules and Indexes")
    bullets(
        doc,
        [
            "users/{uid}: user can read/update own profile fields only. Role fields are not writable.",
            "menuItems/{id}: authenticated users can read. Staff/admin writes use controlled functions or strict rules.",
            "orders/{id}: student can read own order. Staff/admin can read queue/history. Direct creates are denied.",
            "paymentAttempts/{id}: student can read own attempt. Writes are by Cloud Functions only.",
            "stats/today and stats/YYYY-MM-DD: admin read, Cloud Functions write.",
            "counters/{id} and rateLimits/{uid}: no client read/write.",
            "Required index: orders status ASC, createdAt DESC for staff active queue.",
            "Required index: orders businessDate ASC, createdAt DESC for admin history.",
            "Required index: orders userId ASC, createdAt DESC for student history.",
        ],
    )

    h1(doc, 14, "UI and Offline Behaviour")
    bullets(
        doc,
        [
            "Student menu supports cached viewing while offline but blocks order placement with a clear message.",
            "Student tracking uses a single order listener and a confirming-payment fallback screen.",
            "Staff queue displays a prominent connection-lost banner when listener connectivity drops.",
            "Staff UI retains last known queue data and automatically resumes on reconnect.",
            "Staff status buttons show only the next permitted action as primary.",
            "Admin reports distinguish live stats from finalised historical stats.",
        ],
    )

    h1(doc, 15, "Image Upload and Storage")
    bullets(
        doc,
        [
            "Uploader: admin only, enforced by Firebase Storage rules and UI access.",
            "Input types: JPEG and PNG only.",
            "Maximum input size: 5 MB before upload.",
            "Processing: Storage trigger converts to WebP at max 800x600 and writes generated asset path.",
            "Failure handling: menu item remains draft/unpublished until processed image is available, or admin explicitly approves original fallback.",
            "Serving: Next.js Image with explicit dimensions and CDN cache headers.",
        ],
    )

    h1(doc, 16, "Performance and Reliability")
    bullets(
        doc,
        [
            "Menu page uses App Router revalidation, not getStaticProps.",
            "Availability changes call an authenticated revalidation route protected by REVALIDATE_SECRET.",
            "Customer tracking listens only to orders/{orderId}.",
            "Staff queue uses a scoped indexed query for active orders.",
            "Admin dashboard reads stats/today rather than running aggregation queries at page load.",
            "Cloud Functions emit structured logs and alert on payment errors, webhook failures, refund failures, and elevated latency.",
        ],
    )

    h1(doc, 17, "Refund Design")
    bullets(
        doc,
        [
            "cancelOrder initiates a full source refund through Razorpay.",
            "Student-facing copy says: Refund initiated. Timing depends on payment method and bank.",
            "Store refundId, status, speed, createdAt, processedAt, and failureReason if present.",
            "Listen for refund.processed and refund.failed webhooks where supported.",
            "If refund initiation fails, alert admin immediately and preserve an auditable refund_failed state.",
        ],
    )

    h1(doc, 18, "Development Phases")
    bullets(
        doc,
        [
            "Phase 0 - Infrastructure: Firebase projects, Vercel projects, secrets, rules scaffold, indexes, CI. Exit: empty app deployed to dev/staging/prod.",
            "Phase 1 - Auth and roles: phone auth, staff/admin auth, bootstrap script, assignRole. Exit: role tests pass and first admin documented.",
            "Phase 2 - Payment core: initiatePayment, verifyPayment, webhook, idempotency, attempts. Exit: payment and webhook replay tests pass.",
            "Phase 3 - Menu and cart: menu CRUD, ISR revalidation, cart limits, image processing. Exit: availability and image upload tests pass.",
            "Phase 4 - Order flow: student checkout, tracking, cancellation, refunds. Exit: E2E order flow passes on staging.",
            "Phase 5 - Staff operations: live queue, status transitions, offline banner. Exit: queue works with index and reconnect test.",
            "Phase 6 - Admin and reporting: stats/today, stats/YYYY-MM-DD, dashboard, history. Exit: historical stats retained across midnight IST test.",
            "Phase 7 - Hardening: load test, App Check enforcement, alerts, rollback runbook. Exit: production sign-off complete.",
        ],
    )

    h1(doc, 19, "Testing Strategy")
    bullets(
        doc,
        [
            "Cloud Functions unit tests: cart validation, rate limits, HMAC verification, idempotency, cancellation window, refund handling.",
            "Firestore rules tests: role boundaries, denied direct order creation, denied counter writes, admin-only stats.",
            "Webhook tests: valid signature, invalid signature, replay, duplicate captured event, refund processed/failed.",
            "E2E tests: student order, payment success, staff status changes, cancellation, admin menu update, ISR refresh.",
            "Load tests: 200 concurrent users, 50 active orders, queue listener latency, function p95 latency.",
            "Timezone tests: orders and token reset around 23:59 and 00:00 Asia/Kolkata.",
        ],
    )

    h1(doc, 20, "Operations, Monitoring, and Rollback")
    bullets(
        doc,
        [
            "Create alerts for payment verification failures, webhook failures, refund failures, high function latency, Firestore permission-denied spikes, and queue listener errors.",
            "Maintain runbooks for Razorpay outage, Firebase outage, Vercel rollback, and corrupted menu item data.",
            "Production deployments require staging E2E pass, rules test pass, function tests pass, and indexes deployed.",
            "Rollback uses Vercel previous deployment plus Firebase function version rollback where applicable.",
            "Daily export or backup policy must include orders, paymentAttempts, stats, users, and menuItems.",
        ],
    )

    h1(doc, 21, "Known MVP Limitations")
    bullets(
        doc,
        [
            "Single canteen location only.",
            "Availability is binary; no inventory quantity tracking.",
            "No estimated wait time until enough historical preparation data exists.",
            "No loyalty, coupons, table delivery, or pre-order scheduling.",
            "Push notifications beyond in-app/browser-supported status updates are deferred.",
        ],
    )

    h1(doc, 22, "Future Enhancements")
    bullets(
        doc,
        [
            "Multi-canteen support with scoped roles and super-admin reporting.",
            "Inventory stock counts and automatic sold-out toggling.",
            "Preparation-time prediction based on item mix and kitchen load.",
            "Browser push notifications through Firebase Cloud Messaging.",
            "QR table ordering and scheduled pickup slots.",
            "Demand forecasting and procurement planning.",
        ],
    )

    h1(doc, 23, "Open Decisions")
    bullets(
        doc,
        [
            "Maximum order value: owner canteen operator; default recommendation INR 2,000.",
            "Cancellation window: owner canteen operator; default recommendation 90 seconds while placed.",
            "Refund speed: owner finance/admin; default recommendation normal refunds unless instant refund cost is approved.",
            "Admin MFA provider: owner engineering; default recommendation use Firebase-supported MFA path or documented third-party flow.",
            "Original image fallback: owner engineering/admin; default recommendation keep item unpublished if WebP generation fails.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("College Canteen Management System - PDR v3.0 - Confidential")

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
